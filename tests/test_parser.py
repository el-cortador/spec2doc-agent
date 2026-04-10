"""
Тесты модуля core/parser.py.
Используют только реальные in-memory файлы — никаких моков pdfplumber/docx.
"""
import io
import pytest
from docx import Document as DocxDocument
from docx.oxml.ns import qn

# Чтобы pytest нашёл пакет из корня проекта
import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent.parent))

from core.parser import extract_text, ParserError


# ── Вспомогательные фабрики ───────────────────────────────────────────────────

def make_docx_bytes(paragraphs: list[str], table_rows: list[list[str]] | None = None) -> bytes:
    """Создаёт DOCX в памяти с заданными параграфами и, опционально, таблицей."""
    doc = DocxDocument()
    for text in paragraphs:
        doc.add_paragraph(text)
    if table_rows:
        table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for r, row in enumerate(table_rows):
            for c, cell_text in enumerate(row):
                table.cell(r, c).text = cell_text
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Тесты: неподдерживаемый формат ───────────────────────────────────────────

def test_unsupported_extension_raises(tmp_path):
    f = tmp_path / "file.txt"
    f.write_text("hello")
    with pytest.raises(ParserError, match="Неподдерживаемый формат"):
        extract_text(str(f))


# ── Тесты: DOCX ───────────────────────────────────────────────────────────────

def test_docx_paragraphs(tmp_path):
    data = make_docx_bytes(["Первый абзац", "Второй абзац"])
    f = tmp_path / "test.docx"
    f.write_bytes(data)

    result = extract_text(str(f))

    assert "Первый абзац" in result
    assert "Второй абзац" in result


def test_docx_table(tmp_path):
    data = make_docx_bytes([], table_rows=[["Заголовок А", "Заголовок Б"], ["Ячейка 1", "Ячейка 2"]])
    f = tmp_path / "table.docx"
    f.write_bytes(data)

    result = extract_text(str(f))

    assert "Заголовок А" in result
    assert "Ячейка 2" in result


def test_docx_paragraphs_and_table(tmp_path):
    data = make_docx_bytes(
        ["Описание функциональности"],
        table_rows=[["Параметр", "Значение"], ["timeout", "30"]],
    )
    f = tmp_path / "mixed.docx"
    f.write_bytes(data)

    result = extract_text(str(f))

    assert "Описание функциональности" in result
    assert "timeout" in result
    assert "30" in result


def test_docx_empty_raises(tmp_path):
    data = make_docx_bytes([])
    f = tmp_path / "empty.docx"
    f.write_bytes(data)

    with pytest.raises(ParserError, match="не содержит текста"):
        extract_text(str(f))


def test_docx_returns_string(tmp_path):
    data = make_docx_bytes(["Текст"])
    f = tmp_path / "str.docx"
    f.write_bytes(data)

    result = extract_text(str(f))
    assert isinstance(result, str)


# ── Тесты: PDF ────────────────────────────────────────────────────────────────

def test_pdf_no_text_layer_raises(tmp_path):
    """Минимальный валидный PDF без текстового слоя."""
    minimal_pdf = (
        b"%PDF-1.4\n"
        b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
        b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n"
        b"3 0 obj<</Type/Page/MediaBox[0 0 3 3]>>endobj\n"
        b"xref\n0 4\n"
        b"0000000000 65535 f \n"
        b"0000000009 00000 n \n"
        b"0000000058 00000 n \n"
        b"0000000115 00000 n \n"
        b"trailer<</Size 4/Root 1 0 R>>\n"
        b"startxref\n190\n%%EOF"
    )
    f = tmp_path / "no_text.pdf"
    f.write_bytes(minimal_pdf)

    with pytest.raises(ParserError, match="текстовый слой не обнаружен"):
        extract_text(str(f))
